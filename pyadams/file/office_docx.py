"""
    office_docx 操作
        1. 基于python-docx 操作
        2. 基于win32com 操作
"""

# 标准库
import time
import os.path
import logging

# 调用库
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt,RGBColor
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.table import _Cell
from docx.oxml import OxmlElement

# ----------
logger = logging.getLogger('office_docx')
logger.setLevel(logging.DEBUG)
is_debug = True


# ===========================================================================
# win32com 操作
# ===========================================================================

# 合并操作 - 包含图片
def combind_words_win32(main_path, file_paths):
    """
        调用 win32com 模块
    """

    if is_debug: logger.info('combind_words_win32 win32com调用')

    main_path = os.path.abspath(main_path)

    file_paths.reverse() # 顺序调反
    import win32com.client as win32
    #打开word软件
    word = win32.gencache.EnsureDispatch('Word.Application')
    #非可视化运行
    word.Visible = False
    file_paths[0] = os.path.abspath(file_paths[0])
    output = word.Documents.Open(file_paths[0])

    for path in file_paths[1:]:
        path = os.path.abspath(path)
        output.Application.Selection.Range.InsertFile(path)#拼接文档

    #获取合并后文档的内容
    # doc = output.Range(output.Content.Start, output.Content.End)
    # doc.Font.Name = "黑体"  #设置字体

    output.SaveAs(main_path)
    output.Close()

    return None


# docx 转 PDF 
def doc2pdf(doc_path, pdf_path=None):
    """
        调用 win32com 模块
    """
    if is_debug: logger.info('doc2pdf win32com调用')

    doc_path = os.path.abspath(doc_path)

    import win32com.client as win32
    #打开word软件
    try:
        word = win32.Dispatch("Word.Application")
    except:
        word = win32.gencache.EnsureDispatch('Word.Application')
    #非可视化运行
    word.Visible = False
    output = word.Documents.Open(doc_path)
    if pdf_path==None:
        pdf_path = doc_path[:-4]+'pdf'

    pdf_path = os.path.abspath(pdf_path)
    output.SaveAs(pdf_path, FileFormat = 17)

    output.Close()

    return pdf_path


# 文档编辑
class WordEdit:
    """
        指定word文件内柔替换
        wps适用
        $test$ 文字替换
        #TEST# 图片替换
    """
    def __init__(self,word_path):
        
        import win32com.client
        if is_debug: logger.info('WordEdit win32com调用')

        word_path = os.path.abspath(word_path)

        word = win32com.client.Dispatch("Word.Application")
        word.Visible = 0
        word.DisplayAlerts = 0
        document = word.Documents.Open(word_path)
        content = document.Content
        selection = word.Selection

        self.file_path = word_path
        self.word = word
        self.document = document
        self.content = content
        self.selection = selection

    def replace_edit(self,oldList,newList): # 指定替换 文字及图片
        """
            $ 开头为文字替换
            # 开头为图片替换
        """

        for oldStr,newStr in zip(oldList,newList):
            if oldStr[0] == "$":
                self.replace_str(oldStr,newStr)

            elif oldStr[0] == "#":
                self.add_picture(oldStr,newStr)

    def save(self,targetAdr=None): # 另存为
        """ 保存 word"""
        if targetAdr!=None:
            targetAdr = os.path.abspath(targetAdr)
            self.document.SaveAs2(targetAdr,12,False,"",True,"",False,False,False,False,False,15)
        else:
            self.document.SaveAs2(self.file_path,12,False,"",True,"",False,False,False,False,False,15)

        if is_debug: logger.info(f'保存word:{self.file_path}')

    def close(self):
        """ 关闭 word"""
        self.document.Close()
        if is_debug: logger.info(f'关闭word:{self.file_path}')

    def replace_str(self, oldStr, newStr):
        """
            文字替换形式添加
        """
        selection = self.selection

        selection.Find.ClearFormatting()
        selection.Find.Replacement.ClearFormatting()
        selection.Find.Execute(oldStr,False,False,False,False,False,True,1,True,newStr,2)

        if is_debug: logger.info(f'文字替换:{oldStr} --> {newStr}')

    def add_picture(self, oldStr, fig_path):
        """
            图片添加
            以文字替换形式添加
        """
        selection = self.selection
        fig_path = os.path.abspath(fig_path)
        selection.Find.ClearFormatting()
        selection.Find.Replacement.ClearFormatting()
        while True:
            selection.start = 0
            if (selection.Find.Execute(oldStr,False,False,False,False,False,True,1,True,"none",0)):
                selection.Text = ""
                selection.InlineShapes.AddPicture(fig_path,False,True)
            else:
                break

        if is_debug: logger.info(f'图片替换:{oldStr} --> {fig_path}')



# ===========================================================================
# docx 
# ===========================================================================

#合并操作 - 图片无法完成
def combine_words_docx(main_path, file_paths):
    """
        基于 python-docx 模块
    """
    main_doc = Document()

    for index, file in enumerate(file_paths):
        sub_doc = Document(file)

        # Don't add a page break if you've reached the last file.
        if index < len(file_paths)-1:
            sub_doc.add_page_break()

        for element in sub_doc.element.body:
            main_doc.element.body.append(element)

    main_doc.save(main_path)

    return None


# 表格边框设置
def set_cell_border(cell: _Cell, **kwargs):
    """
    Set cell`s border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
 
    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
 
    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
 
            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
 
            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


WIDTH_MAX_CM = 17
FONT_NAME = u'宋体'
FONT_SIZE = 11

# docx文档生成
class DataDocx:

    def __init__(self, docx_path):
        
        if is_debug: logger.info('DataDocx python-docx调用')
        self.document = Document()
        self.docx_path = docx_path
        self.set_page_margin()
        self.set_font()

    def set_page_margin(self, x=1.27, y=1.27):
        # 起始页面宽度设置
        for section in self.document.sections:
            # print('默认页面的宽度和高度：', section.page_width.cm, section.page_height.cm) 
            section.top_margin = Cm(y)
            section.bottom_margin = Cm(y)
            section.left_margin = Cm(x)
            section.right_margin = Cm(x)
        return None

    def set_font(self):
        # 字体默认设置
        self.document.styles['Normal'].font.name = FONT_NAME
        self.document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
        self.document.styles['Normal'].font.size = Pt(FONT_SIZE)
        self.document.styles['Normal'].font.color.rgb = RGBColor(0,0,0)

    def save(self):
        # 保存
        document = self.document
        document.save(self.docx_path)
        # help(document.save)

        return None

    def add_paragraph_center(self, str1):
        # 加中间对齐的行
        p = self.add_paragraph(str1)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        return p

    def add_paragraph(self, str1):
        # 加字段
        p = self.document.add_paragraph(str1)
        paragraph_format = p.paragraph_format
        paragraph_format.space_before=Pt(1)    #上行间距
        paragraph_format.space_after=Pt(1)    #下行间距
        # paragraph_format.line_spacing=Pt(18)  #行距
        return p

    # 导入图片
    def add_docx_figure(self, fig_path, name=None, width=WIDTH_MAX_CM):
        # 导入图片
        document = self.document
        fig_path = os.path.abspath(fig_path)
        fig = document.add_picture(fig_path, width=Cm(width))
        # fig.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        last_paragraph = self.document.paragraphs[-1]
        last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        if name!=None:
            self.add_paragraph_center(name)

        return None

    def add_page_break(self):
        # 另起一页

        return self.document.add_page_break()       

    def add_heading(self, str1, level, size=FONT_SIZE, align='left'):
        # 开头 level 
        h = self.document.add_heading('', level=level)
        run = h.add_run(str1)
        run.font.name = FONT_NAME
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
        run.font.size = Pt(size)
        if align.lower() == 'center':
            h.paragraph_format.alignment =  WD_ALIGN_PARAGRAPH.CENTER
        elif align.lower() == 'left':
            h.paragraph_format.alignment =  WD_ALIGN_PARAGRAPH.LEFT
        elif align.lower() == 'right':
            h.paragraph_format.alignment =  WD_ALIGN_PARAGRAPH.RIGHT

        return h

    # 项目符号 圆符号
    def add_list_bullet(self, str1, size=FONT_SIZE):
        # 项目符号 圆符号
        p = self.document.add_paragraph('', style='List Bullet')
        run = p.add_run(str1)
        run.font.size = Pt(size)

        return p

    # 项目符号 数字符号
    def add_list_number(self, str1, size=FONT_SIZE):
        # 项目符号 数字符号
        p = self.document.add_paragraph('', style='List Number')
        run = p.add_run(str1)
        run.font.size = Pt(size)

        return p

    # 表格
    def add_table(self, table_name, str_list):

        document = self.document
        self.add_paragraph_center('\n'+table_name)

        is_title = True
        table = document.add_table(rows=1, cols=len(str_list[0]))

        for line in str_list:
            if is_title: # 标题
                cells = table.rows[0].cells
                is_title = False
            else:
                cells = table.add_row().cells

            for cell, value in zip(cells, line):
                cell.text = str(value)

                # 边框设置
                set_cell_border(cell,top={"sz": 12, "val": "single", "color": "#000000", "space": "0"},
                    bottom={"sz": 12, "color": "#000000", "val": "single"},
                    start={"sz": 12, "val": "single"},
                    end={"sz": 12, "val": "single"},)

                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                cell.paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                # cell.paragraphs[0].paragraph_format.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        # table.alignment = WD_TABLE_ALIGNMENT.CENTER
        # self.add_paragraph('\n')

        return None

    # 页脚设置-仅字符串输入
    def add_footer_str(self, str1): 
        
        document = self.document

        for section in document.sections:
            footer = section.footer
            paragraph = footer.paragraphs[0]
            paragraph.add_run(str1)

        return None

    # 页眉设置-插入图片
    def add_header_figure(self, fig_path): 

        fig_path = os.path.abspath(fig_path)
        document = self.document

        for section in document.sections:
            header = section.header
            paragraph = header.paragraphs[0]
            run = paragraph.add_run()
            run.add_picture(fig_path, width=Cm(16))
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        return None

    # 封面
    def add_cover_page(self, page_one): 
        """
            封面设置
            page_one = {
                'title_main':'静侧翻稳定性',
                'title_minor':'--ADAMS/Car TILT',
                'file_paths':['a', 'b']
            }
        """

        str_time = time.strftime('%Y.%m.%d %H:%M', time.localtime(time.time()))

        self.add_heading('\n'*2, level=1, size=20)
        self.add_heading(page_one['title_main'], level=1, size=40, align='center')
        self.add_heading(page_one['title_minor'], level=1, size=25, align='right')
        self.add_heading('\n'*1, level=1, size=20)

        for num, file_path in enumerate(page_one['file_paths']):
            self.add_heading(f'File{num}: '+os.path.basename(file_path), level=1, size=15, align='center')
        
        self.add_heading('\n'*(6-num), level=1, size=20)

        self.add_list_bullet('创建时间:'+str_time, size=15)
        self.add_list_bullet('自动生成', size=15)

        return None



# ===========================================================================
# 测试
# ===========================================================================

def test_WordEdit():
    # word文档编辑 测试
    word_path = r'..\code_test\office_docx\word_test.docx'
    oldList = ["$test$", "#TEST#"]
    newList = ["测试---test", r'..\code_test\office_docx\GB_T_s.res_1.png']
    targetAdr = r'..\code_test\office_docx\word_test_new.docx'

    word_obj = WordEdit(word_path)
    word_obj.replace_edit(oldList,newList)
    word_obj.save(targetAdr)
    word_obj.close()


def test_DataDocx():

    docx_path = r'..\tests\file_office_docx\demo4.docx'
    # docx_path = os.path.abspath(docx_path)
    title = 'A / B'
    obj = DataDocx(docx_path)
    obj.add_heading(title, level=0, size=20)
    obj.add_heading('时域数据对比', level=1, size=15)
    obj.add_list_bullet('a vs b 1', size=14)
    obj.add_docx_figure(r'..\tests\file_office_docx\GB_T_s.res_1.png', '图片1')
    obj.add_list_bullet('a vs b 2', size=14)
    obj.add_docx_figure(r'..\tests\file_office_docx\GB_T_s.res_2.png', '图片2')

    # document.add_paragraph('Intense quote', style='Intense Quote')

    str_list = [['Name', '伪损伤', '最大值', '最小值', 'RMS']]*3
    table_name = '\n表格一'
    obj.add_table(table_name, str_list)
    obj.save()

    doc2pdf(docx_path)

    import os
    os.popen(docx_path[:-4]+'pdf')



if __name__ == '__main__':

    pass
    logging.basicConfig(level=logging.INFO)

    # docx_path = r'..\tests\file_office_docx\demo4.docx'
    # # docx_path = os.path.abspath(docx_path)
    # title = 'A / B'
    # obj = DataDocx(docx_path)
    # obj.add_heading(title, level=0, size=20)
    # obj.add_heading('时域数据对比', level=1, size=15)
    # obj.add_list_bullet('a vs b 1', size=14)
    # obj.add_docx_figure(r'..\tests\file_office_docx\GB_T_s.res_1.png', '图片1')
    # obj.add_list_bullet('a vs b 2', size=14)
    # obj.add_docx_figure(r'..\tests\file_office_docx\GB_T_s.res_2.png', '图片2')

    # # document.add_paragraph('Intense quote', style='Intense Quote')

    # str_list = [['Name', '伪损伤', '最大值', '最小值', 'RMS']]*3
    # table_name = '\n表格一'
    # obj.add_table(table_name, str_list)
    # obj.save()

    # doc2pdf(docx_path)


    # import os
    # os.popen(docx_path[:-4]+'pdf')

    # main_path = r'..\code_test\office_docx\main3.docx'
    # docx_path = r'..\code_test\office_docx\demo.docx'
    # docx_path2 = r'..\code_test\office_docx\demo2.docx'
    # main_path = os.path.abspath(main_path)
    # docx_path = os.path.abspath(docx_path)
    # docx_path2 = os.path.abspath(docx_path2)
    # combind_words_win32(main_path, [docx_path, docx_path2])
    # test_WordEdit()